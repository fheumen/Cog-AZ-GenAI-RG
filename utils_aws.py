import re
import os
import csv
import boto3
from datetime import datetime
from const_aws import *
import pdfplumber
import io
from docx import Document as Document_docx
from docx.shared import Inches
import pandas as pd
import tiktoken
import json
from langchain_community.vectorstores import OpenSearchVectorSearch
from opensearchpy import RequestsHttpConnection, OpenSearch
from langchain_community.embeddings import BedrockEmbeddings
from langchain.docstore.document import Document
from dotenv import load_dotenv
from requests_aws4auth import AWS4Auth
from langchain_openai import ChatOpenAI
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from chat_app_aws import ReportGeneration

load_dotenv()

s3 = boto3.client("s3")


0######### Initiate AWS Credentials
service = "aoss"
credentials = boto3.Session().get_credentials()
region = boto3.Session().region_name
awsauth = AWS4Auth(
    credentials.access_key,
    credentials.secret_key,
    region,
    service,
    session_token=credentials.token,
)


##########OpenSearch
opensearch_domain_endpoint = os.environ["OPENSEARCH_ENDPOINT"]
opensearch_index = os.environ["OPENSEARCH_INDEX"]
embedding = BedrockEmbeddings(model_id=os.environ["aws_embedding_model"])

openai_api_key = os.environ.get("OPENAI_API_KEY")

load_dotenv()

chat = ReportGeneration(embedding, opensearch_domain_endpoint, opensearch_index)
#llm = ChatOpenAI(model_name=model_name, temperature=0, openai_api_key=openai_api_key)
llm = chat.chat_model

opensearch_vdb = OpenSearchVectorSearch(
    embedding_function=embedding,
    index_name=opensearch_index,
    http_auth=awsauth,
    use_ssl=True,
    verify_certs=True,
    http_compress=True,  # enables gzip compression for request bodies
    connection_class=RequestsHttpConnection,
    opensearch_url=opensearch_domain_endpoint,
)

# Connect to OpenSearch
opensearch_client = OpenSearch(
    hosts=opensearch_domain_endpoint,
    http_compress=True,
    use_ssl=True,
    verify_certs=True,
    connection_class=RequestsHttpConnection,
    http_auth=awsauth # Replace with your credentials
)
#### Control the size of token
def num_tokens_from_string(string: str, encoding_name: str) -> int:
    #encoding = tiktoken.encoding_for_model(encoding_name)
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens

### Get the list of files into a S3 directory
def get_list_of_files(bucket_name, folder_prefix):
    response = s3.list_objects_v2(Bucket=bucket_name, Prefix=folder_prefix)["Contents"]
    return response


##### Create a Folder into a Bucket
def check_create_dir(bucket_name, path):
    # Check if the subfolder exists by listing objects with the given prefix
    response = s3.list_objects_v2(Bucket=bucket_name, Prefix=path)
    if "Contents" not in response:
        s3.put_object(Bucket=bucket_name, Key=(path + "/"))


#### Read the Mapping column wise
def get_mapping_list(bucket_name, excel_file_path = f"{MAPPING_FILE_PATH}", az_mapping_sheet_name=f"{SHEET_NAME_MAPPING }"):
    #excel_file_path = f"{MAPPING_FILE_PATH}"
    #az_mapping_sheet_name = "Mapping"

    # Get the Excel file from S3
    response = s3.get_object(Bucket=bucket_name, Key=excel_file_path)
    excel_data = response['Body'].read()
    df_mapping = pd.read_excel(io.BytesIO(excel_data), sheet_name=az_mapping_sheet_name)

    section_names_ispr = df_mapping["ISPR Structur Section names"].tolist()
    section_names_keysearch = [x.lower() for x in df_mapping["Section names key search"].tolist()]
    ispr_summary_flag = df_mapping["Integration type in ISPR Flag"].tolist()
    ispr_map_prompt_ls = df_mapping["Map Prompt Template"].tolist()
    ispr_combine_prompt_ls = df_mapping["Combine Prompt Template"].tolist()

    return {
        "section_names_ispr": section_names_ispr,
        "section_names_keysearch": section_names_keysearch,
        "ispr_summary_flag": ispr_summary_flag,
        "ispr_map_prompt_ls": ispr_map_prompt_ls,
        "ispr_combine_prompt_ls": ispr_combine_prompt_ls

    }

# mapping_dic = get_mapping_list(bucket_name, excel_file_path=f"{MAPPING_FILE_PATH}",
#                          az_mapping_sheet_name=f"{SHEET_NAME_MAPPING}")
#
# print(mapping_dic["section_names_keysearch"])
# print(mapping_dic["section_names_ispr"])
# print(mapping_dic["ispr_summary_flag"])
# print(mapping_dic["ispr_map_prompt_ls"])
# print(mapping_dic["ispr_combine_prompt_ls"])


#### Get the Summary of sections from PQRs
def get_abs_summarize(docs,  ispr_map_prompt, ispr_combine_prompt):
    ispr_map_prompt_tmp = PromptTemplate(template=ispr_map_prompt, input_variables=["text"])
    ispr_combine_prompt_prompt_tmp = PromptTemplate(template=ispr_combine_prompt, input_variables=["text"])
    chain = load_summarize_chain(llm, chain_type="map_reduce", map_prompt=ispr_map_prompt_tmp, combine_prompt=ispr_combine_prompt_prompt_tmp )
    #try:
    summary = chain.run(docs)
    #except:
    #docs = docs[:-1]
    return summary

###################### Extract features from the first pages of a PQR
def extract_reporting_period_product_name_site_name(
    first_page_text, date_pattern, site_names, product_names
):

    site_name = None
    product_name = None
    reporting_period = None,
    reporting_period_startdate = None
    reporting_period_enddate = None
    ###### Extract Site Name from the first page
    for site in site_names:
        if site.lower() in first_page_text.lower():
            site_name = site

    ##### Extract Product Name from the first page
    for product in product_names:
        if product.lower() in first_page_text.lower():
            product_name = product

    # Extract Reporting Period from the first page
    match = re.search(date_pattern, first_page_text)
    if match:
        # Check which format of date was matched
        if match.group(1) and match.group(2):
            start_date_str, end_date_str = match.group(1), match.group(2)
            date_format = "%d %b %Y"
        elif match.group(3) and match.group(4):
            start_date_str, end_date_str = match.group(3), match.group(4)
            date_format = "%d %b %Y"
        elif match.group(5) and match.group(6):
            start_date_str, end_date_str = match.group(5), match.group(6)
            date_format = "%B %d, %Y"

            # Convert to datetime objects and format as desired (e.g., 14Nov2022)
        reporting_period_startdate = datetime.strptime(
            start_date_str, date_format
        ).strftime("%d%b%Y")
        reporting_period_enddate = datetime.strptime(
            end_date_str, date_format
        ).strftime("%d%b%Y")

        reporting_period = reporting_period_startdate + "_" + reporting_period_enddate

        # return reporting_period_startdate, reporting_period_enddate
    return reporting_period, reporting_period_startdate, reporting_period_enddate, product_name, site_name


####################
################# Extract features from the first pages of a PQR
def extract_text_tables_images_by_sections(
    pdf_path,
    section_names,
    reporting_period,
    product_name,
    single_filename,
    site_name,
    output_dir,
    chunk_min,
    chunk_max,
):
    # extract_text_tables_images_by_sections(filename_to_proceed, section_names, reporting_period, product_name, pqr_file_name, site_name, output_path_images, 1024, 1500)

    # def extract_text_tables_images_by_sections(pdf_path, section_names, single_filename, output_dir,
    #                                           chunk_min, chunk_max):
    """
    Extract text, tables, and images from a PDF based on a given list of section names.

    Args:
        pdf_path (str): Path to the input PDF file.
        section_names (list): A list of section names to chunk the PDF by.
        output_dir (str): Directory to save extracted images.

    Returns:
        list: A list of dictionaries, each containing 'section_name', 'text', 'images', and 'tables'.
    """

    ####
    header_flag = 0
    footer_flag = 0

    # Normalize section names for matching
    normalized_section_names = [name.lower() for name in section_names]

    # Initialize data structure
    sections = []
    chunk_sections = []
    docs = []  ### list of document object for Vector Database
    current_section = {
        "section_name": None,
        "file_name": single_filename,
        "reporting_period": None,
        "reporting_period_startdate": None,
        "reporting_period_enddate": None,
        "product_name": None,
        "site_name": None,
        "page_num": None,
        "images": [],
        "tables": [],
        "text": "",
    }

    chunk_current_section = {
        "section_name": None,
        "file_name": single_filename,
        "reporting_period": None,
        "reporting_period_startdate": None,
        "reporting_period_enddate": None,
        "product_name": None,
        "site_name": None,
        "page_num": None,
        "images": [],
        "tables": [],
        "text": "",
    }

    # Function to find if a text matches any section name
    def match_section(text):
        # normalized_text = ((text.strip().lower().split(' '))[1]).strip()
        textsplit = text.strip().lower().split(" ")
        if len(textsplit) > 1:
            firstchar = ((text.strip().lower().split(" "))[0]).strip()
            restchar = (" ".join((text.strip().lower().split(" "))[1:])).strip()
            if firstchar.isdigit():
                # return next((name for name in normalized_section_names if restchar.startswith(name)), None)
                return next(
                    (name for name in normalized_section_names if name in restchar),
                    None,
                )
        return None

    # Open the PDF using pdfplumber
    pqr_file_obj = s3.get_object(Bucket=bucket_name, Key=pdf_path)
    pqr_file_data = pqr_file_obj["Body"].read()
    with pdfplumber.open(io.BytesIO(pqr_file_data)) as pdf:
        # with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if (
                page_num == 0
            ):  ##### Extract reporting period, product name and site name from the first page
                (
                    reporting_period,
                    reporting_period_startdate,
                    reporting_period_enddate,
                    product_name,
                    site_name,
                ) = extract_reporting_period_product_name_site_name(
                    text, date_pattern, site_names, product_names
                )

            elif page_num == 1:  ##### Jump the Table of Content
                continue

            # Split text into lines for matching section names
            lines = text.split("\n")
            for line in lines:
                ######### Skip the header and footer
                if re.search(header_pattern, line):
                    header_flag = 1
                    continue
                if header_flag == 1:
                    header_flag = 0
                    continue
                if re.search(footer_pattern, line):
                    footer_flag = 1
                    continue
                if footer_flag == 1:
                    footer_flag = 0
                    continue

                ### No Header, No Footer, No Page 1
                matched_section = match_section(line)
                if matched_section:
                    # Save the previous section
                    if chunk_current_section["section_name"]:
                        doc = Document(
                            page_content=chunk_current_section["text"],
                            metadata=chunk_current_section,
                        )
                        docs.append(doc)
                        chunk_sections.append(chunk_current_section)

                    if current_section["section_name"]:
                        sections.append(current_section)

                    # Start a new section
                    current_section = {
                        "section_name": matched_section,
                        "file_name": single_filename,
                        "reporting_period": reporting_period,
                        "reporting_period_startdate": reporting_period_startdate,
                        "reporting_period_enddate": reporting_period_enddate,
                        "product_name": product_name,
                        "site_name": site_name,
                        "page_num": page_num,
                        "images": [],
                        "tables": [],
                        "text": "",
                    }

                    # Start a new section chunk
                    chunk_current_section = {
                        "section_name": matched_section,
                        "file_name": single_filename,
                        "reporting_period": reporting_period,
                        "reporting_period_startdate": reporting_period_startdate,
                        "reporting_period_enddate": reporting_period_enddate,
                        "product_name": product_name,
                        "site_name": site_name,
                        "page_num": page_num,
                        "images": [],
                        "tables": [],
                        "text": "",
                    }

                # Add the line text to the current section if inside a section
                if current_section["section_name"]:
                    current_section["text"] += line.strip() + "\n"

                if chunk_current_section["section_name"]:
                    # current_section["text"] += line.strip() + "\n"
                    # check if the size of a chunk text into a section is within the chunk size
                    if (
                        len(chunk_current_section["text"] + line.strip()) > chunk_max
                    ) and (len(chunk_current_section["text"]) > chunk_min):
                        # Save a section chunk
                        chunk_current_section_tmp = chunk_current_section.copy()
                        chunk_current_section_tmp["text"] += line.strip() + "\n"
                        ################ create a doc object
                        doc = Document(
                            page_content=chunk_current_section_tmp["text"],
                            metadata=chunk_current_section_tmp,
                        )
                        docs.append(doc)
                        chunk_sections.append(chunk_current_section_tmp)
                        # sections.append(current_section_tmp)

                        chunk_current_section["text"] = ""

                    chunk_current_section["text"] += (
                        line.strip() + "\n"
                    )  # print(sections)

            # Extract tables
            tables = page.extract_tables()
            if tables:
                # for table in tables:

                for table_index, table in enumerate(tables):
                    # Convert the table into a pandas DataFrame
                    df = pd.DataFrame(table)

                    # Create a unique filename for each table (based on page number and table index)
                    sfilename = single_filename.split(".")[0]
                    table_filename = f"{sfilename}_table_page_{page_num + 1}_table_{table_index + 1}.csv"
                    output_csv = f"{output_dir}/{table_filename}"
                    tmp_file = "/tmp/" + table_filename

                    # Save the DataFrame as a CSV file
                    csv_buffer = io.StringIO()
                    df.to_csv(csv_buffer, index=False, header=True)

                    # Upload CSV to S3
                    s3.put_object(
                        Bucket=bucket_name, Key=output_csv, Body=csv_buffer.getvalue()
                    )

                    # table_count += 1
                    # print(f"Table {table_count} saved to {output_csv}")
                    current_section["tables"].append(output_csv)

            # Extract images
            if page.images:
                for img_idx, image_dict in enumerate(page.images):
                    # Extract the image bytes
                    # print("image_dict", image_dict)
                    page_height = page.height
                    image_bbox = (
                        image_dict["x0"],
                        page_height - image_dict["y1"],
                        image_dict["x1"],
                        page_height - image_dict["y0"],
                    )
                    # image = page.within_bbox(image_dict['bbox']).to_image()
                    image = page.within_bbox(image_bbox).to_image()
                    image_bytes = image.original
                    image_ext = "png"  # Can modify based on your preference

                    # Convert image to bytes
                    img_buffer = io.BytesIO()
                    image_bytes.save(img_buffer, format="PNG")
                    img_buffer.seek(0)

                    # Save the image to a file
                    sfilename = single_filename.split(".")[0]
                    image_filename = (
                        f"{sfilename}_page_{page_num + 1}_img_{img_idx + 1}.{image_ext}"
                    )
                    # print("titititititti")
                    # print(image_filename)
                    image_path = f"{output_dir}/{image_filename}"

                    # Upload PNG to S3
                    s3.put_object(Bucket=bucket_name, Key=image_path, Body=img_buffer)

                    # shutil.move(tmp_file, image_path )
                    current_section["images"].append(image_path)

    # Append the last section
    if current_section["section_name"]:
        doc = Document(
            page_content=chunk_current_section["text"],
            metadata=chunk_current_section,
        )
        docs.append(doc)
        sections.append(current_section)
        chunk_sections.append(chunk_current_section)

    return {
        "sections": sections,
        "chunk_sections": chunk_sections,
        "documents": docs,
        "reporting_period": reporting_period,
        "reporting_period_startdate": reporting_period_startdate,
        "reporting_period_enddate": reporting_period_enddate,
        "product_name": product_name,
        "site_name": site_name,
    }


#############


#####
def ingest_pqr_file(tmp_pqr_pdf_path):
    with pdfplumber.open(io.BytesIO(tmp_pqr_pdf_path)) as pdf:
        # Get the first page
        first_page = pdf.pages[0]
        # Extract text from the first page
        first_page_text = first_page.extract_text()
        return extract_reporting_period_product_name_site_name(
            first_page_text, date_pattern, site_names, product_names
        )


def Ingest_phase_1(bucket_name, pqr_file_name):
    tmp_pqr_pdf_path = f"{upload_folder}{pqr_file_name}"
    print(pqr_file_name)

    pqr_file_obj = s3.get_object(Bucket=bucket_name, Key=tmp_pqr_pdf_path)
    pqr_file_data = pqr_file_obj["Body"].read()
    reporting_period, reporting_period_startdate, reporting_period_enddate, product_name, site_name = (
        ingest_pqr_file(pqr_file_data)
    )
    #reporting_period = reporting_period_startdate + "_" + reporting_period_enddate

    #### create the directory structure if not exist: reporting_period/product_name
    # input_folder_pqr_reporting_period = f"{input_folder}/{product_name}"
    input_folder_pqr_reporting_period_product_name = (
        f"{input_folder}{product_name}/{reporting_period}"
    )
    # print(input_folder_pqr_reporting_period_product_name )
    # check_create_dir(input_folder_pqr_reporting_period)
    check_create_dir(bucket_name, input_folder_pqr_reporting_period_product_name)

    #### move file from tmp to reporting_period/product_name
    s3.copy(
        {"Bucket": bucket_name, "Key": tmp_pqr_pdf_path},
        bucket_name,
        f"{input_folder_pqr_reporting_period_product_name}/{pqr_file_name}",
    )
    s3.delete_object(Bucket=bucket_name, Key=tmp_pqr_pdf_path)
    return reporting_period, product_name, site_name


############
from langchain_community.vectorstores import OpenSearchVectorSearch

#### Check if an entry is already available in Pinececone.
def opensearch_document_exists(
     product_name, reporting_period, site_name
):
    # Perform the metadata search
    response = opensearch_client.search(
        index=opensearch_index ,
        body={
            "query": {
                "bool": {
                    "must": [
                        {"term": {"metadata.product_name.keyword": product_name}},
                        {"term": {"metadata.reporting_period.keyword": reporting_period}},
                        {"term": {"metadata.site_name.keyword": site_name}}
                    ]
                }
            }
        }
    )

    # Check if any results were returned
    if response['hits']['total']['value'] > 0:
        print("Metadata already exists:", response['hits']['hits'])
        return 1
    else:
        print("Metadata does not exist.")
        return 0

def opensearch_insert_docs(documents):

    print(f"Going to insert {len(documents)} to Pinecone")
    # Prepare embeddings and documents for insertion
    opensearch_vdb.add_documents(
        documents=documents,
    )
    print("****** Added to Pinecone vectorstore vector")
    ################


def Ingest_phase_2(product_name, reporting_period, site_name, pqr_file_name):
    filename_to_proceed = (
        f"{input_folder}{product_name}/{reporting_period}/{pqr_file_name}"
    )
    output_path_images_tables = f"{input_folder}{product_name}/{reporting_period}"
    #### if the doc link (pdf, docx, txt, doc) was already saved in Pinecone, then ignore, otherwise insert.
    # if (pinecone_document_noexists(pqr_file_name, reporting_period, product_name, site_name, INDEX_NAME)== 1):
    if pqr_file_name.endswith(".pdf"):
        sec_doc = extract_text_tables_images_by_sections(
            filename_to_proceed,
            section_names,
            reporting_period,
            product_name,
            pqr_file_name,
            site_name,
            output_path_images_tables,
            1024,
            1500,
        )
        return sec_doc

        # section_chunks = section_chunks  + sec_doc["sections"]

        # else:
        #     if single_filename.endswith(".docx"):
        #        section_chunks = section_chunks + (ingest_docx(filename_to_proceed, section_names, reporting_period, product_name, single_filename, site_name))
        # pinecone_insert_docs(sec_doc["documents"], INDEX_NAME)
        # opensearch_insert_docs(sec_doc["documents"])


def save_chunks_to_json(chunks, json_filename):
    """Save the chunked sections to a JSON file."""

    # tmp_file = "/tmp/" + json_filename
    output_path = f"{output_folder}{json_filename}"

    # Convert the list of dictionaries to a JSON string
    json_str = json.dumps(chunks, indent=4)

    # Convert the JSON string to a bytes object
    json_bytes = io.BytesIO(json_str.encode("utf-8"))

    s3.upload_fileobj(json_bytes, bucket_name, output_path)


def Ingest_PQR(bucket_name, upload_folder):
    list_of_upload_files = get_list_of_files(bucket_name, upload_folder)
    sections = []
    section_chunks = []
    if len(list_of_upload_files) > 0:
        for pqr_file in list_of_upload_files:
            pqr_file_key = pqr_file["Key"]
            if pqr_file_key.endswith(".pdf") or pqr_file_key.endswith(".docx"):
                pqr_file_name = pqr_file_key.split("/")[-1]
                reporting_period, product_name, site_name = Ingest_phase_1(
                    bucket_name, pqr_file_name
                )
                sec_doc = Ingest_phase_2(
                    product_name, reporting_period, site_name, pqr_file_name
                )

                sections = sections + sec_doc["sections"]
                section_chunks = section_chunks + sec_doc["chunk_sections"]
                # print(sec_doc["documents"])

                # pinecone_insert_docs(sec_doc["documents"], INDEX_NAME)
                if opensearch_document_exists(product_name, reporting_period, site_name) == 0:
                    opensearch_insert_docs(sec_doc["documents"])
                else:
                    print("Metadata already exists:")

    ispr_json_filename = "ispr" + "-" + product_name + "-" + reporting_period + ".json"
    ispr_chunk_json_filename = (
        "ispr_chunk" + "-" + product_name + "-" + reporting_period + ".json"
    )

    save_chunks_to_json(sections, ispr_json_filename)
    save_chunks_to_json(section_chunks, ispr_chunk_json_filename)

### TAking the json file produced by Ingest_pqr() and creating an ispr docx file
def ispr_generation(bucket_name):

    mapping_dic = get_mapping_list(bucket_name, excel_file_path=f"{MAPPING_FILE_PATH}",
                         az_mapping_sheet_name=f"{SHEET_NAME_MAPPING}")

    section_names_keysearch = mapping_dic["section_names_keysearch"]
    section_names_ispr = mapping_dic["section_names_ispr"]
    ispr_summary_flag = mapping_dic["ispr_summary_flag"]
    ispr_map_prompt_templates = mapping_dic["ispr_map_prompt_ls"]
    ispr_combine_prompt_templates = mapping_dic["ispr_combine_prompt_ls"]


    for product_name in product_names:

        for reporting_period in reporting_periods:

            ispr_json_filename = "ispr" + "-" + product_name + "-" + reporting_period + ".json"
            output_path_json = f"{output_folder}{ispr_json_filename}"
            json_obj = s3.get_object(Bucket=bucket_name, Key=output_path_json)
            json_data = json_obj['Body'].read().decode('utf-8')  # Read and decode the body

            df = pd.DataFrame(json.loads(json_data))
            document = Document_docx()

            for section_index, section_name in enumerate(section_names_keysearch):
                i = 1
                document.add_heading(section_names_ispr[section_index], i)
                i = i + 1

                #### Collecting the text, images and tables present in a specific section
                docs = []
                images_section_sites = []
                tables_section_sites = []

                for site_name in site_names:

                    # Subsetting specific rows and columns by labels
                    subset = df.loc[
                        (df["section_name"] == section_name.lower())
                        & (df["site_name"] == site_name.lower())
                        & (df["product_name"] == product_name)
                        & (df["reporting_period"] == reporting_period),
                        ["page_num", "images", "tables", "section_name", "text"],
                    ]

                    if not subset.empty:
                        max_row = subset.loc[subset["page_num"].idxmax()]
                        if ispr_summary_flag[section_index] == 1:
                            text = max_row["text"]
                            for i in range(0, len(text), model_max_tokens):
                                chunk = text[i:i + model_max_tokens]
                                doc = Document(
                                     page_content=chunk,
                            # metadata=result["matches"][j]["metadata"],
                                )
                                docs.append(doc)
                            images_section_sites.append(max_row["images"])
                            tables_section_sites.append(max_row["tables"])


                        else:
                            heading_name = section_name + "_" + site_name
                            document.add_heading(heading_name, i)
                            document.add_paragraph(max_row["text"])
                        # Add images
                            for image_path in max_row["images"]:
                            #print(image_path)
                               if len(image_path) > 0:
                                  img_stream = io.BytesIO()
                                  s3.download_fileobj(bucket_name, image_path, img_stream)
                                  img_stream.seek(0)

                                  document.add_picture(
                                    img_stream, width=Inches(4)
                                )  # Add image with a fixed width

                        # Add tables
                            for table_path in max_row["tables"]:
                               if len(table_path) > 0:

                                   csv_stream = io.BytesIO()
                                   s3.download_fileobj(bucket_name, table_path, csv_stream)
                                   csv_stream.seek(0)

                                  # Read the CSV content
                                   csv_data = list(csv.reader(io.StringIO(csv_stream.read().decode('utf-8'))))

                                  # Create a new table in Word
                                   rows = len(csv_data)
                                   cols = len(csv_data[0])
                                   word_table = document.add_table(rows=rows, cols=cols)

                                  # Populate the Word table
                                   for row_index, row in enumerate(csv_data):
                                       for col_index, cell in enumerate(row):
                                           word_table.cell(row_index, col_index).text = str(cell)

                    else:
                        continue
                if ispr_summary_flag[section_index] == 1:
                    summary_text = get_abs_summarize(docs, ispr_map_prompt_templates[section_index], ispr_combine_prompt_templates[section_index])
                    document.add_paragraph(summary_text)

            ispr_filename = "ispr" + "_" + product_name  + "_" + reporting_period + ".docx"
            target_filename = f"{output_folder}{ispr_filename}"

            # Save the document to a BytesIO object (in-memory file)
            doc_stream = io.BytesIO()
            document.save(doc_stream)
            doc_stream.seek(0)
            s3.upload_fileobj(doc_stream, bucket_name, target_filename)

