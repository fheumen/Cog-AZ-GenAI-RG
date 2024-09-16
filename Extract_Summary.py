# Databricks notebook source
import os
import zipfile
import re
import xml.dom.minidom
import pandas as pd
import docx
import logging
import logging.config
from docx2python import docx2python
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Length, Pt
from copy import deepcopy
import glob
import os
import csv
from docx.oxml.ns import qn
from openpyxl import workbook #pip install openpyxl
from openpyxl import load_workbook
from datetime import datetime

# COMMAND ----------

def IS_BeginExtraction (para, title_begin = "summary and conclusion"):   
    
    if len(re.findall(title_begin.lower(), para.text.lower())) > 0: 
         
        ########### "Summary and conclusion" als Titel
        tst = para.text.split('\t') 
        if len(tst) == 3:                      
            return False
        else: return True
    return False       

# COMMAND ----------

def IS_StopExtraction(para, title_end = "batches reviewed "):  
    #BATCHES REVIEWED (APPROVED AND REJECTED)
    #batches reviewed (approved and rejected) 
    
    if len(re.findall(title_end.lower(), para.text.lower())) > 0:
        if(('Heading' in  para.style.name) or  ("Überschrift" in  para.style.name) or ("titre" in  para.style.name.lower()) \
       or ("title" in  para.style.name.lower()) or ("level" in  para.style.name.lower()) or ("gliederung" in  para.style.name.lower()) \
       or re.search("^\d+\.*\d*\s+\w+", para.text) is not None): 
           return True
    
    else: return False
        

# COMMAND ----------

def del_FooterandHeader(doc):     
    ##### Remove Headers and Footers in Documents
    for section in doc.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True

# COMMAND ----------

def del_ActiveFields(body_elements): 
    #body_elements = doc._body._body
    ps = body_elements.xpath('.//w:fldChar')
    for i in range(len(ps)):
        ps[i].getparent().remove(ps[i])
    
    ps = body_elements.xpath('.//w:instrText')
    for i in range(len(ps)):
         ps[i].getparent().remove(ps[i]) 

    ps = body_elements.xpath(".//w:hyperlink")
    label_r = body_elements.xpath('.//w:hyperlink/w:r')
    for i in range(len(ps)):        
        ps[i].addprevious(label_r[i])
        ps[i].getparent().remove(ps[i])   

# COMMAND ----------

def del_userdefinedcontainer(body_elements):  
    #body_elements = doc._body._body
    ps = body_elements.xpath('.//w:sdt')
    for i in range(len(ps)):
         ps[i].getparent().remove(ps[i])

# COMMAND ----------

def format_BodyText(paragraph, font_name = 'Times New Roman'):
    for run in paragraph.runs:
        font = run.font
        #font.size = Pt(font_size_text)
        font.name = font_name 
        ####

# COMMAND ----------

def format_TableText(table, font_name = 'Times New Roman'):    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = font_name 
                    #font.size= Pt(font_size_table)

# COMMAND ----------

def check_of_succ_iSRS(abs_filename, max_heading = 5):
    doc = docx.Document(abs_filename)
    cnt_title = 0
    for para in doc.paragraphs:   
    
       if IS_StopExtraction(para):
        #or  ("Überschrift" in  para.style.name) or ("titre" in  para.style.name.lower()) \
          #or ("title" in  para.style.name.lower()) or ("level" in  para.style.name.lower()) or ("gliederung" in  para.style.name.lower()): 
            cnt_title = cnt_title + 1
            if cnt_title > max_heading: return False   
            
    return True

# COMMAND ----------

def define_log_file(status_iSRS, filename):
    
    LOG_FILENAME = status_iSRS + filename + ".log"
    tmp_log_file = "/tmp/"+LOG_FILENAME

    logging.config.dictConfig({
        'version': 1,
        'disable_existing_loggers': True,
    })

    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    logging.basicConfig(format= '%(asctime)s — %(name)s — %(levelname)s — %(message)s', filemode='w', filename=tmp_log_file, level=logging.INFO)   
    
    logging.info('Extraction Job Started...')
    
    return (LOG_FILENAME)
    

# COMMAND ----------

def iSRS_write_in_log(res_code, filename, abs_filename = ""):
    
    if res_code == 0: 
        LOG_FILENAME = define_log_file("error", filename)
        logging.error("   " + filename.ljust(105) + "  -->" + '  Failed: There is no "SUMMARY" in this Document')
        logging.info('Extraction Job Ended...')
        
    elif res_code == 2:
        LOG_FILENAME = define_log_file("error", filename)
        logging.warning(" " + filename.ljust(105) + "  -->" + '   Failed: There is no "Conclusion" in "SUMMARY" in this Document') 
        logging.info('Extraction Job Ended...')

    elif res_code == 3: 
        LOG_FILENAME = define_log_file("error", filename)
        logging.error("   " + filename.ljust(105) + "  -->" + '  Failed: Unable to open the file. Probably the file was not saved properly.')
        logging.info('Extraction Job Ended...')
        
    elif res_code == -1 or res_code == -2 or res_code == -5: 
        if check_of_succ_iSRS(abs_filename):
                LOG_FILENAME = define_log_file("", filename)
                if res_code == -5:   logging.info("   " + filename.ljust(105) + "  -->" + '  Pass: Successful "SUMMARY" Extraction, but a Table in SUMMARY have not been formatted as expected')
                else: logging.info("    " + filename.ljust(105) + "  -->" + '  Pass: Successful "SUMMARY" Extraction')
                logging.info('Extraction Job Ended...')
        else:
            LOG_FILENAME = define_log_file("warning", filename)
            logging.warning(" " + filename.ljust(105) + "  -->" + '  Warning: the extracted summary has probably more pages as expected')
            logging.info('Extraction Job Ended...')
    else:
        LOG_FILENAME = define_log_file("warning", filename)
        err_desc = '  Pass: Stopp Condition have not been met'    
        logging.warning(" "+filename.ljust(105) + "  -->" + err_desc)
        logging.info('Extraction Job Ended...')
        
    return (LOG_FILENAME)

# COMMAND ----------

def Extract_Summary(parent, title_begin = "summary and conclusion" , title_end = "batches reviewed", font_name = 'Times New Roman', conclusion_flag = 0): 

    #### Initalise flag variable, helping to identify of the "Text" to extract into the document
    #global flag 
    
    flag = 0
    flg = 0
    
    #### Get the content of all the docx document: Text Paragraphs + Tables
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    ##### Going through the entire Document page per page   
    for child in parent_elm.iterchildren():      
    
        
        ################### 1. If the Text is a Paragraph  ###################            
        if isinstance(child, CT_P):
            para = Paragraph(child, parent)
            
            
            #p = para._p
            #hyperlink = p.xpath("./w:hyperlink")
            #label_r = p.xpath("./w:hyperlink/w:r")
            #hyperlink.add_previous(label_r)
            #p.remove(hyperlink)
            
            if len(re.findall(title_end, para.text.lower())) > 0:
                print(para.style.name)
            
            ##### Identification of the beginning of the text to be extracted:
            if IS_BeginExtraction(para, title_begin) and flag == 0:  
                print(para.text)
                flag = 1 
            
            ##### Identification of the body of the text to be extracted till the key Word "Conclusion"
            elif flag == 1: flag = 2                
            elif conclusion_flag == 1 and (len(re.findall("Conclusion", para.text)) > 0 \
                                        or len(re.findall("CONCLUSION", para.text)) > 0) and  flag == 2:                    
                    flag = 3
            #print(para.style.name)
            #elif conclusion_flag == 0:
                   # flag = 3
            #elif flag == 3 and IS_StopExtraction(para) == False:
                   # print(para.text)
                    #print(para.style.name)
                    
            ##### Identification of the Stop of extraction
            elif IS_StopExtraction(para, title_end) and  flag == 3: 
                    flag = -1

            elif flag ==2 and IS_StopExtraction(para, title_end) and  conclusion_flag == 0: 
                    print(para.text)
                    flag = -1
            
           
            

            
            
           ############## Delete paragraphs before and after the Text to be extracted ############     
            if flag == 0 or flag == -1 or flag == 1:
                 #delete_paragraph(para)
                para._element.getparent().remove(para._element)
            
            else:########### Formatierung des Body Text in Summary
                format_BodyText(para, font_name = 'Times New Roman')
                    
                
                
                        
        ################## 2. If the Text is a Table #################
        elif isinstance(child, CT_Tbl): 
            table = Table(child, parent)
             ############## Delete Tables before and after the Text to be extracted ############  
            if flag == 0 or flag == -1 :
                table._element.getparent().remove(table._element)
                
            else: ########### Formatierung des Table Text in Summary
                ##print("vvvvvvvvvvvvvvvvvvvvvvvv")
                try:
                    format_TableText(table, font_name = 'Times New Roman')
                except:
                    flag = -5
                    pass
                
    
    ######## In case "Summary" is the last section
    if flag == 3: flag = -2        
    if conclusion_flag == 0 and flag == 2: flag = -2
    
    
    ##### Remove Headers and Footers in Document
    del_FooterandHeader(parent)
    #for section in parent.sections:
    #        section.different_first_page_header_footer = False
    #        section.header.is_linked_to_previous = True
    
    #########################################
    body_elements = parent._body._body
    
    #### Remove "sdt" container in documents 
    del_userdefinedcontainer(body_elements)
    #ps = body_elements.xpath('.//w:sdt')
    #for i in range(len(ps)):
    #     ps[i].getparent().remove(ps[i])
            
    #### Remove "bookmark/Hyperlinks/active Fields" container in documents
    del_ActiveFields(body_elements)
    #ps = body_elements.xpath('.//w:fldChar')
    #for i in range(len(ps)):
         #ps[i].getparent().remove(ps[i])
    
    #ps = body_elements.xpath('.//w:instrText')
    #for i in range(len(ps)):
         #ps[i].getparent().remove(ps[i])
    
        
    return flag
