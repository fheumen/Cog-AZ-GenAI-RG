# Databricks notebook source
from docx import Document
import pandas as pd
from const import *
import shutil
import logging
import logging.config
import docx
import os
from docx.shared import Inches

# COMMAND ----------

output_folder = f"{OUTPUTS_PATH}"  ### output directory, where all output files are save
output_path_json = f"{output_folder}/pqr_section_chunks.json"
# Read the JSON file into a DataFrame

# COMMAND ----------

df = pd.read_json(output_path_json)
df["site_name"] = df["site_name"].str.strip().str.lower()
# Display the DataFrame
# subset_test = df.loc[(df['section_name'] == "other") & (df['site_name'] == "fmc"), ['page_num', 'section_name', 'content']]

subset_test = df.loc[(df['section_name'] == "summary and conclusion") & (df['site_name'] == "fmc") & (df['product_name'] == "Fasenra") & (df['product_name'] == "Fasenra") & (df['reporting_period'] == "14Nov2022_13Nov2013"), ['page_num',  'images']]

max_row_test = subset_test.loc[subset_test['page_num'].idxmax()]

# Display the row with the maximum value
#print(max_row_test)
#print(max_row_test['text'])

print(subset_test)

# COMMAND ----------


#product_names = ["Lumoxity", "Enhertu", "Vaxzevria", "Beyfortus", "IMJUDO", "IMFINZI", "Synagis", "Saphnelo", "Fasenra", "Tezspire"]
for reporting_period in reporting_periods:

    for product_name in product_names:
        document = Document()
        for section_name in section_names:
            i = 1
            document.add_heading(section_name, i)
            i = i + 1
            for site_name in site_names:
               
            # Subsetting specific rows and columns by labels    
                subset = df.loc[(df['section_name'] == section_name.lower()) & (df['site_name'] == site_name.lower()) & (df['product_name'] == product_name) & (df['reporting_period'] == reporting_period), ['page_num', 'section_name', 'text']]
                print(subset)
                print(subset.empty)
                print(not subset.empty)
                if not subset.empty:
                   print(subset)
                   max_row = subset.loc[subset['page_num'].idxmax()]
                   heading_name = section_name + "_" + site_name
                   document.add_heading(heading_name, i)

                   document.add_paragraph(max_row['text'])
                    # Add images
                   for image_path in df["images"]:
                       if len(image_path) > 0:
                           print(image_path)
                           document.add_picture(image_path[0], width=Inches(4))  # Add image with a fixed width
        
                    # Add tables
                   for table in df["tables"]:
                       if len(table) > 0:
                           # Create a new table in Word
                          rows = len(table)
                          cols = len(table[0])
                          word_table = document.add_table(rows=rows, cols=cols)
                
                    # Populate the Word table
                   for row_index, row in enumerate(table):
                       for col_index, cell in enumerate(row):
                           word_table.cell(row_index, col_index).text = str(cell)

                else: continue
                   
        ispr_filename = "ispr" + "_"  + reporting_period + "_" + product_name + ".docx"
        tmp_target_file = "/tmp/"+ispr_filename
        target_filename =f"{output_folder}/{ispr_filename}"
        
        document.save(tmp_target_file)                
        shutil.move(tmp_target_file,target_filename)

    
