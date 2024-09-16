# Databricks notebook source
# MAGIC %run ./Utils/StorageAccountAccess_Secrets

# COMMAND ----------

# MAGIC %run ./Extract_Summary

# COMMAND ----------

import os
import shutil
from dotenv import load_dotenv
load_dotenv()

# COMMAND ----------

#secret_scope = "nlp-toolbox_storage_account"
#basepath=mount_blob_storage(secret_scope)

storageAccountName = os.environ.get("storageAccountName")
#storageAccountAccessKey = os.environ.get("storageAccountAccessKey")
storageAccountAccessKey = "KqDmutvYKFkUfibvOLivxyuYP+x9KqY3DblOGarDbgJjMBEbEfUn3bhEjyuNO+DUvgDdOcl2fgrX+AStRyyYDA=="
print(storageAccountAccessKey)
mount_blob_storage("inputs")
mount_blob_storage("outputs")

#input_folder =dbutils.secrets.get(secret_scope,"input_folder")
#output_folder = dbutils.secrets.get(secret_scope,"output_folder")

input_folder = f"/dbfs/mnt/FileStore/inputs/"  ### intput directory, where all intputs  files are save
output_folder = f"/dbfs/mnt/FileStore/outputs/"  ### output directory, where all output files are save


# COMMAND ----------

dbutils.fs.mounts()

# COMMAND ----------

#dbutils.fs.unmount(dbutils.fs.mounts()[0].mountPoint)

# COMMAND ----------

dbutils.fs.ls("/mnt/FileStore")

# COMMAND ----------

import docx
doc = docx.Document("dbfs:/mnt/FileStore/inputs/Sample_ISPR.docx")

# COMMAND ----------


list_of_files = get_list_of_files('/dbfs/mnt/FileStore/inputs/')
print(list_of_files)

# COMMAND ----------

import shutil
import logging
import logging.config
import docx
import os

list_of_files = get_list_of_files(input_folder)

if len(list_of_files) > 0:
######## iSRS Extration of docx documents
    #logging.info('Extraction Job Started...')
    ###
    for single_filename in list_of_files:
        print(single_filename)
        if(single_filename.endswith(".docx")):        
            filename_to_proceed = input_folder+single_filename
            print(filename_to_proceed)
            try:
               doc = docx.Document(filename_to_proceed)
            except ValueError:
               res_code = 3
               LOG_FILENAME = iSRS_write_in_log(res_code, single_filename)
               os.remove(filename_to_proceed)
               tmp_log_file = "/tmp/"+LOG_FILENAME
               shutil.move(tmp_log_file, f"{output_folder}/{LOG_FILENAME}")
               continue

            res_code = Extract_Summary(doc, title_begin="REPROCESSED & REWORKED BATCHES", title_end= "PRODUCT REVIEWS FROM PREVIOUS MANUFACTURING STAGE") 
            print(res_code)
            
            if res_code == 2: ##### No Conclusion in Summary
                doc = docx.Document(filename_to_proceed)
                res_code = Extract_Summary(doc, title_begin="REPROCESSED & REWORKED BATCHES", title_end= "PRODUCT REVIEWS FROM PREVIOUS MANUFACTURING STAGE", conclusion_flag = 0)
            
            if res_code not in [0, 2]:
                tmp_target_file = "/tmp/"+single_filename
                target_filename =f"{output_folder}/{single_filename}"
        
                doc.save(tmp_target_file)
                LOG_FILENAME = iSRS_write_in_log(res_code, single_filename, tmp_target_file)
                shutil.move(tmp_target_file,target_filename)
                
            else: LOG_FILENAME = iSRS_write_in_log(res_code, single_filename)
            #Must be enabled in productive !!
            os.remove(filename_to_proceed)
        
    #logging.info('Extraction Job Ended...')
            tmp_log_file = "/tmp/"+LOG_FILENAME
            shutil.move(tmp_log_file, f"{output_folder}/{LOG_FILENAME}")

# COMMAND ----------


